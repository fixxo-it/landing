"""One-off: builds the FamCare content-comparison doc for the founder review.
Not part of the app — run manually, not imported anywhere."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x01, 0x61, 0x63)
INK = RGBColor(0x0B, 0x1F, 0x20)
MUTED = RGBColor(0x5A, 0x6B, 0x6C)

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = INK

def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_title():
    h = doc.add_heading(level=0)
    run = h.add_run("FamCare — Website Copy Comparison")
    run.font.color.rgb = TEAL
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub = doc.add_paragraph()
    r = sub.add_run("Live site (famcare.co.in)  vs.  New design (famcare-landing.vercel.app)")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED
    r.italic = True

    note = doc.add_paragraph()
    r2 = note.add_run(
        "Extracted verbatim from both live pages for founder review. Add / edit / mark up directly in this doc."
    )
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MUTED
    doc.add_paragraph()

def add_section(title, rows, col_headers=("Current site (famcare.co.in)", "New design (vercel)")):
    h = doc.add_heading(level=2)
    r = h.add_run(title)
    r.font.color.rgb = TEAL

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, text in enumerate(col_headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], "016163")

    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

add_title()

add_section("Header / Navigation", [
    ("Title tag: \u201cFamCare \u2013 In-House Trained & Verified Baby Caregivers in Whitefield, Bangalore\u201d", "\u2014"),
    ("Nav: Services | Safe360\u2122 | About us | How it works | Download Now", "Nav: Services | How it works | Safety 360\u00b0 | FAQs | BOOK NOW"),
    ("Badge: \u201cLive now \u00b7 Whitefield, Bangalore\u201d", "\u2014"),
])

add_section("Hero", [
    ("\u201cVerified BabyCare, at your door in 10 min.\u201d", "\u201cOn demand baby care you can trust\u201d"),
    ("\u201cPolice verified \u00b7 In-house trained \u00b7 Arrives in ~10 min\u201d", "(folded into subhead)"),
    ("\u201cIn-house trained, authorised caregivers \u2014 criminal record & background verified. Safety built into every booking, for both you and your caregiver.\u201d",
     "\u201cProfessionally trained, background-verified caregivers at your door in 10 minutes.\u201d"),
    ("Buttons: \u201cDownload Now\u201d / \u201cSafety Features \u2192\u201d", "Buttons: \u201cFIND NEARBY CAREGIVER\u201d / \u201cSEE HOW IT WORKS\u201d"),
    ("\u201cAlso coming soon: Elderly Care\u201d", "\u2014"),
    ("Stats strip: ~10 min avg arrival \u00b7 6-Layer verification \u00b7 \u20b90 starting price \u00b7 6AM\u20139PM service hours \u00b7 0% book-again rate "
     "(note: some stats render as \u201c0 min\u201d / \u201c0-layer\u201d on the live site \u2014 likely a bug, flag before reusing)", "No stats strip"),
])

add_section("Services", [
    ("\u201cCare for every stage\u201d", "\u201cFrom newborn to school going, we have the right caregiver\u201d"),
    ("\u201cFrom newborn to school-going \u2014 we have the right caregiver.\u201d", ""),
    ("\u201cEvery caregiver is matched to the specific needs of your child's age group.\u201d", ""),
    ("Newborn care \u2014 \u201cFeeding support, soothing, and nap routines\u201d", "Newborn care (0\u20136mo)"),
    ("Infant day care \u2014 \u201cSafe at-home supervision, bottle prep, and active caregiving for 3\u201312 months\u201d (badge: \u201cMost booked\u201d)", "Infant day care (6\u201318mo)"),
    ("Toddler companion \u2014 \u201cPlay-based engagement, meals, and routine monitoring for 1\u20133 years\u201d", "Toddler companion (1.5\u20134yr)"),
    ("After-school babysitting \u2014 \u201cDrop-off support, snack time, and supervised homework for 4\u201310 years\u201d", "After school care (4\u201310yr)"),
    ("Elderly care (Coming Soon) \u2014 \u201cCompassionate at-home support, mobility assistance, and daily routine help for seniors\u201d",
     "Elderly care (60+yr) \u2014 live, not marked \u201ccoming soon\u201d"),
])

add_section("Safety / Trust System", [
    ("\u201cFamCare Safe360\u2122\u201d / \u201cEvery booking. 360\u00b0 protected.\u201d", "\u201cSafety 360\u00b0\u201d"),
    ("\u201cA complete safety ecosystem that protects your family before, during and after every booking.\u201d", "\u2014"),
    ("Stats: 6 Safety pillars \u00b7 100% Verified caregivers \u00b7 Live Always monitored", "\u2014"),
    ("Button: \u201cBook a caregiver\u201d", "\u2014"),
    ("Six pillars: Identity360\u00b0 \u00b7 Arrival360\u00b0 \u00b7 Care360\u00b0 \u00b7 Monitor360\u00b0 \u00b7 Quality360\u00b0 \u00b7 Support360\u00b0",
     "Four pillars: Trust, Care, Technology, Quality (see below)"),
    ("\u201cTrust infrastructure\u201d / \u201cFour systems working as one.\u201d / \u201cSafe360 combines people, technology, care protocols and quality controls around every booking.\u201d / \u201cOne connected safety system\u2014from verification to continuous improvement.\u201d",
     "(condensed into the 4-pillar block)"),
    ("Trust \u2014 \u201cIdentity checks, police verification and behaviour scores.\u201d", "Trust \u2014 \u201cIdentity checks, police verification and behaviour scores.\u201d (same wording)"),
    ("Technology \u2014 \u201cAI risk detection, live tracking and safety escalation.\u201d", "Technology \u2014 \u201cRisk detection, live tracking and instant escalation.\u201d (drops \u201cAI\u201d)"),
    ("Care \u2014 \u201cNo-phone baby mode, care logs and family support.\u201d", "Care \u2014 \u201cNo-phone baby mode, care logs and a family support line.\u201d"),
    ("Quality \u2014 \u201cFeedback loops, audits and continuous certification.\u201d", "Quality \u2014 \u201cFeedback after every visit, regular audits, ongoing certification.\u201d"),
])

add_section("About FamCare", [
    ("\u201cAbout FamCare\u201d / \u201cBuilding the operating system for trust in home care.\u201d", "not present"),
    ("\u201cFamCare is building India's AI-enabled technology infrastructure for caregiving\u2014so families no longer have to rely on word of mouth and guesswork.\u201d", ""),
    ("\u201cWe operate as a technology company, not simply a caregiving marketplace. Verified and FamCare Academy-trained caregivers are supported by in-session monitoring, geofencing, and a dedicated Command Center that watches over every session in real time.\u201d", ""),
    ("\u201cBy building industrial-grade trust at scale, we help working couples and working women pursue their careers without compromising their families' safety.\u201d", ""),
    ("Link: \u201cExplore FamCare Safe360\u2122 \u2192\u201d", ""),
])

add_section("How It Works", [
    ("\u201cHow FamCare works\u201d (Active system)", "\u201cBook in 3 taps\u201d"),
    ("FamCare Academy \u2014 \u201cIn-house caregiver training and care readiness.\u201d", "Screen 1: \u201cWhen do you need care?\u201d \u2014 Instant / Schedule toggle, nearest caregivers list"),
    ("Identity & trust \u2014 \u201cBackground, police, and face verification.\u201d", "Screen 2 (service selection, per phone demo)"),
    ("Session intelligence \u2014 \u201cAudio signals, geofencing, and smart alerts.\u201d", "Screen 3: \u201cKavya is on the way\u201d"),
    ("Command Center \u2014 \u201cReal-time operations and incident response.\u201d", ""),
    ("\u201cHuman oversight across every layer\u201d", ""),
    ("Booking-process block: \u201cSimple & fast\u201d / \u201cFrom need to caregiver in 10 minutes.\u201d / \u201cNo agency calls, no 3-day wait, no negotiations.\u201d "
     "Steps: 01 Open the app (30 sec) \u2192 02 Caregiver is assigned (1 min) \u2192 03 Confirm & pay (2 min) \u2192 04 Caregiver arrives (~10 min)",
     "(merged into the 3-tap phone demo)"),
])

add_section("Verification", [
    ("\u201cZero compromise\u201d / \u201cEvery caregiver is authorised, verified & trained.\u201d", "\u201cNever alone\u201d / \u201cCleared before they knock\u201d"),
    ("\u201cWe don't just take anyone. Our 6-layer process means you know exactly who is walking into your home. Safety is non-negotiable \u2014 for our users and our caregivers.\u201d",
     "\u201cSix independent checks stand between an application and your front door, and they are re-run while the caregiver is with us.\u201d"),
    ("Criminal Records & Background Verification Check \u2014 \u201cVerified against national criminal records and police databases. Renewed every 6 months.\u201d", "\u201cCriminal and police records\u201d"),
    ("Government ID verification \u2014 \u201cAadhaar and PAN cross-checked and stored securely.\u201d", "\u201cGovernment ID, cross-checked\u201d"),
    ("Employment history check \u2014 \u201cPast employers contacted and references verified.\u201d", "\u201cEmployment history and references\u201d"),
    ("Skill assessment \u2014 \u201cPractical baby care skills tested by our in-house trainers.\u201d", "\u201cPractical skill assessment\u201d"),
    ("In-person interview \u2014 \u201cEvery caregiver meets our ops team before going live.\u201d", "\u201cIn-person interview\u201d"),
    ("Physical address verified \u2014 \u201cHome address confirmed and on record before any assignment.\u201d", "\u201cHome address confirmed\u201d"),
    ("\u2b50 Ongoing rating review \u2014 \u201cDrop below 4.2 stars? Immediately paused and retrained.\u201d",
     "\u201cRe-run every six months. Anyone who slips below 4.2 stars is paused and retrained.\u201d"),
    ("Sample caregivers: Priya S. (Newborn specialist, 4yrs, 4.9\u2605, 312 tasks) / Sunita K. (Toddler companion, 5yrs, 4.8\u2605, 274 tasks) / Radha M. (Infant day care, 3yrs, 4.9\u2605, 198 tasks)",
     "not present \u2014 generic \u201cKavya S.\u201d etc. used as filler names elsewhere"),
    ("\u201cReady to book a verified caregiver?\u201d / Button: \u201cBook Now \u2192\u201d", "not present as a standalone block"),
])

add_section("Live Monitoring / Oversight", [
    ("(folded into tech features)", "\u201cEyes on the visit\u201d \u2014 \u201cKnow where your caregiver is, and look in on your child whenever you want, without having to ask.\u201d"),
    ("", "\u201cVisit in progress \u00b7 42 min\u201d / \u201cLook in on the room\u201d \u2014 \u201cLive video, any time you want\u201d"),
    ("", "\u201cSomeone is always watching over it\u201d \u2014 \u201cEvery visit is watched over while it happens. When something needs attention, a trained person is already on it.\u201d"),
    ("", "\u201cReal people, real time\u201d / \u201cOn watch\u201d \u2014 Live oversight, Early warning, Fast escalation, Location safety"),
    ("", "\u201cRisks are surfaced early, escalated quickly, and closed by a real person.\u201d"),
])

add_section("Emergency / SOS", [
    ("\u201c1-tap emergency\u201d / SOS Button \u2014 \u201cEmergency help, instantly\u201d", "\u201cHelp in one tap\u201d \u2014 \u201cOne press reaches our response team and your emergency contacts at the same moment.\u201d"),
    ("\u201cOne tap sends an immediate alert to our response team and your emergency contacts. Available to both parents and caregivers \u2014 because safety can't wait.\u201d",
     "\u201cSOS\u201d / \u201cOne tap, both at once\u201d \u2014 \u201cOur response team\u201d / \u201cYour emergency contacts\u201d"),
    ("", "\u201cOpen to parents and caregivers both, on every visit.\u201d"),
])

add_section("Caregiver Hiring & Training", [
    ("(not a distinct section \u2014 folded into \u201cZero compromise\u201d verification block)", "\u201cSelected for character. Trained for care.\u201d"),
    ("", "1. Careful hiring \u2014 \u201cWe shortlist for temperament first \u2014 patience, calm communication, cleanliness and empathy \u2014 across two interview rounds.\u201d"),
    ("", "2. Verified identity \u2014 \u201cGovernment ID, address and criminal background verification are completed before a caregiver is ever shown to a family.\u201d"),
    ("", "3. Mandatory training \u2014 \u201cClassroom and hands-on training in baby care, hygiene, emergency response and parent communication.\u201d"),
    ("", "4. Practical assessment \u2014 \u201cSkills are demonstrated, observed and certified in-centre. Only then does a caregiver become bookable.\u201d (Score: 94)"),
])

add_section("Technology Features", [
    ("\u201cBuilt for trust\u201d / \u201cOur tech keeps your child safe at all times.\u201d",
     "(distributed across the Trust/Care/Tech/Quality pillars and the oversight section rather than a standalone feature grid)"),
    ("Live Tracking \u2014 \u201cAlways know where your caregiver is\u201d / \u201cTrack your caregiver in real time from the moment they leave until they arrive. Get live updates, estimated arrival time, and peace of mind \u2014 every step of the way.\u201d", ""),
    ("Video Calling \u2014 \u201cCheck in on your little one anytime\u201d / \u201cConnect face-to-face with your caregiver at any moment during the session. See your baby, ask questions, and stay involved \u2014 without being in the room.\u201d", ""),
    ("SOS Button \u2014 \u201cEmergency help, instantly\u201d (as above)", ""),
    ("Live Audio Recording \u2014 \u201cFull session documentation\u201d / \u201cEvery session is recorded with consent. Review interactions, document milestones, and keep a complete record of your child's care \u2014 always protected and encrypted.\u201d", ""),
    ("Geo-Fencing Alerts \u2014 \u201cReal-Time Safety Monitoring\u201d / \u201cIf a caregiver exits the designated care zone during a session, instant alerts are sent to both parents and the FamCare field audit team \u2014 helping ensure your child's safety at all times.\u201d", ""),
    ("Badges: End-to-end encrypted \u00b7 Location never stored \u00b7 \ufe0f DPDP compliant \u00b7 \u26a1 SOS response in under 90 sec", ""),
])

add_section("Flexible Hours / Situations (current site only)", [
    ("\u201cFlexible hours\u201d / \u201cPay for only the hours you need.\u201d \u2014 1 hour (\u201cQuick help\u201d) \u00b7 3 hours (\u201cQuick support\u201d) \u00b7 5 hours (Coming Soon, \u201cExtended care\u201d)", "not present"),
    ("\u201cFor every situation\u201d / \u201cWhatever comes up, we're ready.\u201d \u2014 Post-delivery support / Date night / outing / Sick child support / Work from home", "not present"),
])

add_section("Comparison Table (current site only)", [
    ("\u201cWhy switch\u201d / \u201cThe old way of finding a nanny is broken.\u201d / \u201cSee how FamCare compares to the traditional agency or referral model.\u201d "
     "\u2014 full 11-row comparison table (time to find, background verification, police check, ID verification, address verification, SOS, geo-fencing, "
     "live tracking, pricing, backup, reviews) \u2014 Agency/referral vs FamCare", "not present"),
])

add_section("Testimonials", [
    ("\u201cReal families, real reviews\u201d / \u201cThey trusted FamCare. Now they won't go back.\u201d / \u201cAll reviews are from verified app store users in Whitefield, Bangalore.\u201d",
     "\u201cWhat are families saying about FamCare\u201d"),
    ("18 real testimonials, all 5\u2605, attributed with first name + surname/initial, \u201cVerified App Review \u00b7 [Month] 2026\u201d \u2014 Puja Baranwal, Siwani Dubey, Nikhil, "
     "Aahan Senapati, Gayatri Panda, Siddhartha Mahapatra, Neha (\u00d73), Archana Kammar, \u0c85\u0cb0\u0ccd\u0c9a\u0ca8\u0cbe KK, Dravisha, Samhithaa Saravanan, Swapna, "
     "Gyanvi Gupta, Suryansh, \u0cae\u0cbe\u0cb0\u0cbf\u0caf\u0cbe Kinkhabwala, Mayank",
     "18 testimonials, attributed with full name + neighbourhood + family role (e.g. \u201cWhitefield \u00b7 Mom of 2\u201d) \u2014 Puja Baranwal (quote rewritten, "
     "same \u201cGreen Wali Didi\u201d detail), Sneha Rao, Arjun Menon, Divya Krishnan, Rahul Iyer, Ananya Shetty, Meghana Reddy, Karthik Nair, Priya Deshpande, "
     "Nikhil Bhat, Lakshmi Venkat, Sara Thomas, Tanvi Joshi, Vikram Rao, Shruti Kulkarni, Aditya Pillai, Neha Agarwal, Ravi Subramaniam"),
    ("Full quotes are the real, unpolished app-store reviews \u2014 short, plain language",
     "\u26a0 Full quotes are longer, composed placeholder copy \u2014 NOT real reviews. Only \u201cPuja Baranwal\u201d overlaps with a real name/story from the live site."),
])

add_section("Final CTA", [
    ("\u201cGo. We've got them.\u201d", "\u201cAt your door in 10 minutes\u201d"),
    ("\u201cYour baby is in safe, verified hands. Book in 2 minutes, caregiver at your door in 10 minutes.\u201d", "\u201cYour child's favorite companion, right when you need them most.\u201d"),
    ("Button: \u201cDownload Now \u2192\u201d", "Button: \u201cBOOK NOW\u201d"),
])

add_section("FAQ", [
    ("\u201cGot questions?\u201d / \u201cFrequently asked questions\u201d / \u201cEverything you need to know about safety, training, bookings, and more.\u201d",
     "\u201cHave questions? We got answers\u201d"),
    ("Tabs: \ufe0f Trust & Safety | Caregiver Training | Booking & Scheduling | \u200dCaregivers | Pricing & Payments | Service Area",
     "Tabs: Trust & safety | Training | Booking | Caregivers | Pricing & payments | Service area (same six categories)"),
    ("Q: \u201cHow are FamCare caregivers verified?\u201d \u2014 \u201cEvery FamCare caregiver goes through a rigorous multi-step verification process before becoming "
     "bookable. This includes application screening, two rounds of interviews, Aadhaar and identity verification, address verification, criminal background "
     "checks, reference checks, mandatory classroom and practical training, and a comprehensive skill assessment. Only candidates who successfully complete "
     "every stage are certified and deployed to serve families.\u201d",
     "Same Q, near-identical answer: \u201cEvery caregiver goes through a multi-step verification process before becoming bookable: application screening, two "
     "rounds of interviews, Aadhaar and identity verification, address verification, criminal background checks, reference checks, mandatory classroom and "
     "practical training, and a comprehensive skill assessment. Only candidates who clear every stage are certified and deployed to families.\u201d"),
    ("Q: \u201cIs my baby safe with a FamCare caregiver?\u201d (answer not extracted)", "Q: \u201cCan I see who is coming before they arrive?\u201d (answer not extracted this pass)"),
    ("Q: \u201cWhat happens if there's an issue during a booking?\u201d (answer not extracted)", "Q: \u201cWhat happens in an emergency?\u201d (answer not extracted this pass)"),
    ("\u201cStill have questions?\u201d contact prompt (implied)", "\u201cStill have questions?\u201d / support@famcare.co.in"),
])

add_section("Footer", [
    ("Logo + \u201cTrusted Baby Care in 10 mins\u201d", "\u201cFamCare\u201d + \u201cTrusted Baby care in 10 minutes\u201d"),
    ("\u201cWhitefield - Bangalore\u201d", "\u201cWhitefield, BLR\u201d"),
    ("+91 95357 11078 \u00b7 support@famcare.co.in", "+91 95357 11078 \u00b7 support@famcare.co.in"),
    ("Services: Newborn care, Infant day care, Toddler companion, After-school care", "Services: Newborn care, Infant day care, Toddler companion, After school care"),
    ("Company: Safe360\u2122, About FamCare, How it works, Support", "Company: Safe 360, About us, How it works, Support"),
    ("Legal: Privacy Policy, Terms & Conditions, Refund Policy, Data Deletion", "Legal: Privacy policy, Terms & Conditions, Refund policy, Data deletion"),
    ("\u00a9 2026 FAMCARE TECHNOLOGIES PRIVATE LIMITED. All rights reserved. Made with \u2764 for Whitefield, Bangalore families",
     "\u00a9 2026 FamCare. All rights reserved. Made with \u2665 for families in Bengaluru."),
])

# closing notes
h = doc.add_heading(level=2)
r = h.add_run("Notes for review")
r.font.color.rgb = TEAL
for line in [
    "Testimonials on the new site are placeholder copy, not the real app-store reviews from the live site \u2014 only \u201cPuja Baranwal\u201d reuses a real "
    "detail (the \u201cGreen Wali Didi\u201d line). These need replacing with real reviews before launch, or a clear placeholder flag.",
    "Sections that exist only on the live site, with no equivalent yet on the new build: About FamCare, Flexible Hours, Situations (\u201cFor every "
    "situation\u201d), the Agency-vs-FamCare comparison table, and the standalone Technology Features grid (Video Calling, Live Audio Recording specifically).",
    "The live site's stats strip appears to have a rendering bug (\u201c0 min\u201d, \u201c0-layer\u201d, \u201c0%\u201d) \u2014 verify the real site before carrying "
    "those numbers over.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(line)

out_path = r"C:\Users\Chinmay Shinde\Downloads\FamCare-Copy-Comparison.docx"
doc.save(out_path)
print("saved:", out_path)
